#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gen_rapor3.py — Proje 3 akademik rapor taslagini (.docx) uretir.
Bicim: Courier New 10pt (PDF gereksinimi). 7 bolum + PC eslemesi + tablolar.
Eksik/placeholder kalan yerler "EKSIK / YAPILMASI GEREKEN" kutulariyla isaretlenir.
Cikti: SISTEMPROGRAMLAMA-2/rapor3_taslak.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Courier New"
SIZE = 10

doc = Document()

# ── Stil: tüm belge Courier New 10pt ────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(SIZE)


def _set_font(run, size=SIZE, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def H(text, size=12, color=(0x1F, 0x4E, 0x79), align=None, space_before=10):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_font(r, size=size, bold=True, color=color)
    return p


def P(text, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    _set_font(r, bold=bold)
    return p


def BUL(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_font(r)
    return p


def CODE(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(12)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        _set_font(r, size=9)
    return p


def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.paragraphs[0].clear()
        _shade_cell(c, "D9E1F2")
        r = c.paragraphs[0].add_run(h)
        _set_font(r, size=9, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].clear()
            r = cells[j].paragraphs[0].add_run(str(val))
            _set_font(r, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def TODO(text):
    """Eksik/yapilmasi gereken kismi sari kutu + kirmizi metinle isaretler."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shade_cell(cell, "FFF2CC")  # acik amber
    cell.paragraphs[0].clear()
    r1 = cell.paragraphs[0].add_run("» EKSİK / YAPILMASI GEREKEN: ")
    _set_font(r1, size=9, bold=True, color=(0xC0, 0x00, 0x00))
    r2 = cell.paragraphs[0].add_run(text)
    _set_font(r2, size=9, color=(0xC0, 0x00, 0x00))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ════════════════════════════════════════════════════════════════════════
# KAPAK
# ════════════════════════════════════════════════════════════════════════
P("T.C. SAKARYA UYGULAMALI BİLİMLER ÜNİVERSİTESİ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P("SİSTEM PROGRAMLAMA DERSİ 3. PROJE ÖDEVİ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P("", space_after=2)
P("PICORV İŞLEMCİ ALT KÜMESİ (RV32I) İÇİN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P("FPGA TABANLI LOADER TASARIMI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
P("TAKIM ÜYELERİ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
for m in ["- ALİHAN HOYLU", "- FEYZA ÖZDEN", "- ENES DURAN", "- MUHAMMED EREN KENDİR"]:
    P(m, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
P("", space_after=2)
P("2026 BAHAR YARIYILI", align=WD_ALIGN_PARAGRAPH.CENTER)
P("(Rapor yazım dili: Courier New, 10 pt — akademik rapor)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
TODO("Bu bir TASLAKTIR. Sarı kutular tamamlanması gereken yerleri gösterir. "
     "Teslimden önce: dosyayı PDF'e çevirip resmi adla kaydedin "
     "(BIL302_PROJE3_A.XX_B.YY_C.ZZ_170526.PDF), kapağa öğrenci numaralarını ekleyin.")

# ── PÇ eşleme tablosu ────────────────────────────────────────────────────
H("PROGRAM ÇIKTILARI (PÇ) EŞLEME", size=12)
P("Bu rapordaki çalışmalar aşağıdaki program çıktıları (PÇ) temelinde "
  "yapılandırılmıştır. İlgili PÇ, her bölümün başında parantez içinde belirtilmiştir.")
TABLE(
    ["Bölüm", "İçerik", "PÇ"],
    [
        ["1. Giriş ve Literatür", "Loader/haberleşme/CRC literatürü", "PÇ6, PÇ13"],
        ["2. Sistem Mimarisi (Co-Design)", "Toolchain arayüzü, bellek haritası, FSM", "PÇ6"],
        ["3. Deneysel Çalışmalar", "Test senaryoları, yükleme süresi, kaynak", "PÇ7"],
        ["4. Küresel/Toplumsal/Ekonomik", "Sürdürülebilirlik, güvenlik, e-atık (SKA)", "PÇ8"],
        ["5. Proje Yönetimi", "RACI, koordinasyon, sürüm kontrol", "PÇ12, PÇ13"],
        ["6. Bireysel Katkı Beyanı", "Her üyenin özgün katkısı", "PÇ12"],
        ["Loader Doğruluğu", "CRC'li kararlı loader", "PÇ1"],
    ],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# BÖLÜM 1
# ════════════════════════════════════════════════════════════════════════
H("BÖLÜM 1. GİRİŞ VE LİTERATÜR ARAŞTIRMASI (PÇ6, PÇ13)", size=13)
P("Bir bilgisayar programının kaynak kodundan donanım üzerinde fiziksel "
  "yürütülmesine kadar uzanan araç zinciri (toolchain); assembler, linker ve "
  "loader olmak üzere ardışık katmanlardan oluşur. Önceki iki projede bu zincirin "
  "ilk iki halkası gerçeklenmiştir: 1. projede RV32I komut setini makine koduna "
  "çeviren iki geçişli bir assembler, 2. projede ise bağımsız nesne dosyalarını "
  "birleştiren bir linker tasarlanmış ve üretilen .mem dosyası FPGA bitstream'ine "
  "SENTEZ ZAMANINDA ($readmemh ile) gömülmüştür. Bu yaklaşımda her yeni programın "
  "çalıştırılması, FPGA'in baştan sentezlenmesini ve yeniden programlanmasını "
  "gerektirir.")
P("Bu proje kapsamında zincirin son halkası olan LOADER (yükleyici) tasarlanmıştır. "
  "Geliştirilen sistem, programı ÇALIŞMA ZAMANINDA seri haberleşme (UART) üzerinden "
  "FPGA'e canlı olarak yükler; böylece yeniden sentez ihtiyacı ortadan kalkar. "
  "Bilgisayar tarafındaki Python betiği, linker'dan çıkan .bin imajını CRC-16 ile "
  "korunan paketler hâlinde gönderir; FPGA üzerindeki sonlu durum makinesi (FSM) "
  "tabanlı loader bu paketleri doğrular, blok belleğe (BRAM) yazar ve yükleme "
  "bitince PicoRV32 işlemcisinin reset hattını serbest bırakarak programı fiziksel "
  "olarak çalıştırır. (PÇ6)")

H("1.1. Gömülü Sistemlerde Program Yükleme Mimarileri", size=11)
P("Gömülü sistemlerde uygulama kodunun işlemci belleğine yüklenmesi için "
  "literatürde üç temel arayüz öne çıkmaktadır:")
BUL("UART tabanlı yükleme: Düşük pin sayısı ve basitliği nedeniyle mikrodenetleyici "
    "bootloader'larında en yaygın yöntemdir. AVR (STK500), STM32 sistem bootloader'ı "
    "ve ESP32 ROM bootloader'ı UART üzerinden firmware yükler [2][5].")
BUL("SPI tabanlı yükleme: Daha yüksek veri hızı sunar; çoğu FPGA, konfigürasyon "
    "bitstream'ini harici SPI flash'tan okur. Senkron saat hattı gerektirir.")
BUL("JTAG tabanlı yükleme: Hata ayıklama ve sınır-tarama (boundary scan) standardı "
    "olan IEEE 1149.1 üzerinden bellek/yazmaç erişimi sağlar; geliştirme aşamasında "
    "yaygındır ancak protokolü karmaşıktır.")
P("Bu projede; düşük donanım maliyeti, tek yönlü sade veri akışı ve Tang Nano 9K "
  "üzerindeki yerleşik USB-UART köprüsüyle doğrudan uyumu nedeniyle UART seçilmiştir. "
  "Seçilen yaklaşım, bir mikrodenetleyici bootloader'ının FPGA üzerinde yumuşak "
  "çekirdek (soft-core) PicoRV32 için yeniden uyarlanmış bir örneğidir [4][5].")

H("1.2. Seri Haberleşme ve Veri Doğrulama Protokolleri", size=11)
P("Seri hatlar üzerinde gürültü, saat kayması (clock drift) veya tampon taşması "
  "nedeniyle bit hataları oluşabilir. Makine kodunun yüklenmesinde tek bir bit "
  "hatası bile programın yanlış çalışmasına yol açacağından, veri bütünlüğünün "
  "doğrulanması zorunludur. Literatürde iki temel yöntem incelenmiştir:")
BUL("Checksum (toplama sağlaması): Baytların modülo toplamı. Hesabı ucuzdur ancak "
    "bit yer değiştirmelerine ve çift hatalara karşı zayıftır.")
BUL("CRC (Cyclic Redundancy Check): Veriyi bir polinoma bölerek kalan üretir. "
    "Burst (ardışık) hataları yüksek olasılıkla yakalar; donanımda kaydırma-XOR ile "
    "ucuza gerçeklenir [3][6].")
P("Bu projede CRC-16/CCITT-FALSE (üretici polinom 0x1021, başlangıç 0xFFFF) "
  "seçilmiştir. Bu polinom 16 bit'e kadar tüm burst hatalarını ve çoğu çoklu hatayı "
  "yakalar; ITU-T V.41 ve XMODEM gibi standartlarda kullanılır [3][6]. Aynı CRC "
  "hesabı host (Python, crc16.py) ve FPGA (Verilog, loader.v içindeki crc_next "
  "fonksiyonu) tarafında BİREBİR gerçeklenmiş ve standart '123456789' vektörü ile "
  "0x29B1 sonucu üzerinden doğrulanmıştır. CRC tutmayan paket NACK ile reddedilir ve "
  "host paketi yeniden gönderir.")

H("1.3. RISC-V / PicoRV32 ve FPGA Hedefi", size=11)
P("Hedef işlemci, açık kaynaklı RISC-V RV32I komut setini uygulayan, boyut "
  "optimizasyonlu PicoRV32 yumuşak çekirdeğidir [4]. Çekirdek, basit bir bellek "
  "veri yolu (mem_valid/mem_ready el sıkışması) ile BRAM'e bağlanır ve PROGADDR_RESET "
  "parametresiyle belirlenen adresten (0x00000000) komut getirmeye başlar [1][4]. "
  "Fiziksel hedef, Gowin GW1NR-9 FPGA içeren Tang Nano 9K kartıdır.")
P("Literatür kriteri: Bu rapor en az üç akademik/standart kaynağa atıf yapar — "
  "RISC-V resmî ISA dokümanı [1], gömülü loader literatürü [2][5], CRC/veri bütünlüğü "
  "standardı [3][6] ve PicoRV32 referans tasarımı [4]. (Kaynakça için bkz. Bölüm 7.)")

# ════════════════════════════════════════════════════════════════════════
# BÖLÜM 2
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
H("BÖLÜM 2. SİSTEM MİMARİSİ VE DONANIM-YAZILIM ORTAK TASARIMI (CO-DESIGN) (PÇ6)", size=13)
P("Sistem, biri bilgisayarda (yazılım) diğeri FPGA'de (donanım) çalışan iki "
  "iş birlikli alt sistemden oluşur. Bu donanım-yazılım ortak tasarımı (co-design), "
  "aynı paket protokolü ve aynı CRC algoritmasının her iki tarafta da gerçeklenmesi "
  "ile sağlanmıştır.")
CODE(
"  BILGISAYAR (HOST)            UART 115200 8N1            TANG NANO 9K (FPGA)\n"
"  +-------------------+   AA 55 01 ADDR N DATA CRC -->  +---------------------+\n"
"  | loader_host.py    |   <-------- ACK / NACK -------- | uart_rx -> loader   |\n"
"  |  .bin -> paket    |                                 |   (CRC-16) -> BRAM  |\n"
"  |  CRC + ACK bekle  |                                 |   cpu_run -> PicoRV |\n"
"  |  RUN -> dinle     |   <------- sonuc bayti -------   |   -> LED / uart_tx  |\n"
"  +-------------------+                                 +---------------------+")
TODO("Bu ASCII blok şema yerine, rapor kalitesi için Word/draw.io ile çizilmiş "
     "GERÇEK bir blok diyagram (şekil) ekleyin ve 'Şekil 2.1' olarak numaralandırın.")

H("2.1. Toolchain Arayüz Standartları", size=11)
P("Araç zinciri, önceki projelerden devralınan özgün arayüzler üzerine kuruludur:")
BUL(".asm -> Assembler: RV32I kaynak kodu iki geçişli assembler ile makine koduna çevrilir.")
BUL("ObjectFile -> Linker: Linker, text/data section'larını text_base=0x0'dan itibaren "
    "yerleştirir ve relocation'ları çözer.")
BUL("Linker -> .bin: Bu projede linker'a eklenen get_bin_output() metodu, linked_code "
    "listesini en küçük adresten itibaren LITTLE-ENDIAN düz ikili imaja çevirir "
    "(PicoRV32 little-endian olduğu için). Ayrıca get_words_hex() ile her satırda bir "
    "32-bit word içeren .words.hex dosyası üretilir (testbench ve insan okuması için).")
BUL("Paket protokolü (host <-> FPGA): Aşağıda 2.2'de tanımlanan WRITE/RUN çerçeveleri "
    "ortak arayüz sözleşmesidir.")
P("Paket protokolü (loader.v ve loader_host.py'de birebir):")
CODE(
"  WRITE: AA 55 01  ADDR[31:0](BE,4B)  NWORDS(1B)  DATA(4*N, BE word)  CRC_HI CRC_LO\n"
"  RUN  : AA 55 02  CRC_HI CRC_LO\n"
"  Cevap: ACK=0x06   NACK=0x15\n"
"  CRC  : CRC-16/CCITT-FALSE (0x1021, init 0xFFFF), [CMD..son veri bayti] uzerinden")

H("2.2. FPGA Loader ve PicoRV32 Bellek Haritası (Memory Map)", size=11)
P("CPU'nun gördüğü bellek haritası aşağıdaki gibidir. BRAM, yükleme anında loader "
  "tarafından, çalışma anında CPU tarafından kullanılır; bu paylaşım top_loader.v "
  "içindeki bir port multiplekseri ile sağlanır.")
TABLE(
    ["Adres aralığı", "Bölge", "Açıklama"],
    [
        ["0x00000000 - 0x00001FFF", "BRAM (8 KB)", "text + data + stack (sp=0x1FF0)"],
        ["0x02000000", "GPIO LED", "yazma: bit[5:0]=LED5..LED0 (aktif-düşük)"],
        ["0x03000000", "UART veri", "yazma: bit[7:0] gönder; okuma: bit0=tx_busy"],
    ],
)
P("Loader FSM durumları (loader.v): Yükleme süresince cpu_run=0 ile CPU reset'te "
  "tutulur; RUN doğrulanınca cpu_run=1 yapılır. Aşağıdaki durum geçiş tablosu FSM'i "
  "özetler:")
TABLE(
    ["Durum", "İşlev", "Sonraki"],
    [
        ["S_SYNC0/1", "AA 55 senkron baytlarını bekle", "S_CMD"],
        ["S_CMD", "Komutu oku, CRC'yi 0xFFFF'ten başlat", "S_A3 / S_RCRC1"],
        ["S_A3..S_A0", "4 baytlık hedef adresi al (BE)", "S_NW"],
        ["S_NW", "Word sayısını al", "S_DATA / S_CRC1"],
        ["S_DATA", "4 bayt topla -> word -> BRAM'e yaz", "S_CRC1 (bitince)"],
        ["S_CRC1/0", "Gelen CRC'yi al ve karşılaştır", "S_REPLY"],
        ["S_REPLY", "ACK/NACK gönder", "S_SYNC0"],
        ["S_RCRC1/0", "RUN paketi CRC kontrolü", "S_RREPLY"],
        ["S_RREPLY", "ACK gönder", "S_RFLUSH"],
        ["S_RFLUSH", "ACK bitince cpu_run=1", "S_RUNNING"],
        ["S_RUNNING", "CPU çalışır, loader bekler", "-"],
    ],
)
TODO("Yukarıdaki durum tablosunun yanına, dairesel durumlar ve geçiş oklarından "
     "oluşan GÖRSEL bir FSM diyagramı (şekil) ekleyin ('Şekil 2.2'). draw.io ile "
     "kolayca çizilir; rubrikte 'şekil' istendiği için kalite puanını yükseltir.")
P("Önemli tasarım kararı — reset hattı yönetimi: RUN paketinin ACK'i, UART TX hattı "
  "henüz loader'a aitken (cpu_run hâlâ 0) gönderilir; ACK gönderimi tamamen bitene "
  "kadar (S_RFLUSH) beklenir, ANCAK ONDAN SONRA cpu_run=1 yapılır. Böylece hem ACK "
  "host'a ulaşır hem de UART TX hattı CPU'ya geçince çalışan program kendi sonucunu "
  "gönderebilir. CPU reset'i ayrıca 3 saatlik bir senkronizer ile temizlenir.")
P("BRAM port multiplekseri (top_loader.v): loading=~cpu_run iken yazma portunu loader "
  "(mem_we/mem_waddr/mem_wdata) sürer; çalışma anında CPU'nun veri yolu bağlanır. Bu "
  "sayede tek portlu BSRAM hem yükleme hem yürütme için paylaşılır.")

# ════════════════════════════════════════════════════════════════════════
# BÖLÜM 3
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
H("BÖLÜM 3. DENEYSEL ÇALIŞMALAR, TEST VE ANALİZ (PÇ7)", size=13)
P("Sistemin yalnızca çalıştığı değil, FARKLI ZORLUK seviyelerinde KARARLI çalıştığı "
  "sistematik bir deney metodolojisi ile gösterilmiştir. Doğrulama, donanım "
  "gerektirmeyen, tekrarlanabilir bir Icarus Verilog testbench'i (tb_loader.v) ile "
  "uçtan uca yapılmıştır: host(TB) -> uart_rx -> loader -> BRAM -> CPU -> uart_tx -> host(TB).")

H("3.1. Deney Tasarımı ve Test Senaryoları", size=11)
P("Test programları, kasıtlı olarak artan karmaşıklıkta seçilmiştir (basit "
  "matematikten fonksiyon çağrılarına). Her program, sonucunu hem 6 LED'e (ikilik) "
  "hem de UART üzerinden host'a gönderir; bu iki bağımsız gözlem kanalı sonucu "
  "çapraz doğrular.")
TABLE(
    ["Test", "Senaryo / Zorluk", "Hesap", "Kullanılan özellik"],
    [
        ["test1_math", "Düz akış (dallanma yok)", "6 + 7 = 13", "ALU, lui, sw/lw"],
        ["test2_loop", "Şartlı döngü", "1+2+...+8 = 36", "bne, geri dallanma"],
        ["test3_func", "Fonksiyon + yığın + iç içe çağrı", "fib(10) = 55", "jal/jalr, sp, callee-save"],
        ["tb_nack", "Bozuk CRC -> NACK -> yeniden gönderim", "13 (test1)", "hata kontrolü"],
    ],
)
P("Yükleme, paket başına 8 word olacak şekilde parçalanmıştır; böylece test3 (38 word) "
  "BEŞ ayrı WRITE paketiyle yüklenir ve çoklu-paket + paket-başı CRC mekanizması "
  "zorlanır.")

H("3.2. Veri Toplama ve Donanım Metrikleri", size=11)
P("Test sonuçları (simülasyon, tb_loader.v / tb_nack.v):")
TABLE(
    ["Test", "Paket sayısı", "UART sonucu", "LED (ikilik)", "Durum"],
    [
        ["test1_math", "2 WRITE + RUN", "13", "001101", "GEÇTİ"],
        ["test2_loop", "2 WRITE + RUN", "36", "100100", "GEÇTİ"],
        ["test3_func", "5 WRITE + RUN", "55", "110111", "GEÇTİ"],
        ["tb_nack", "1 bozuk -> NACK, 2 doğru", "13", "001101", "GEÇTİ"],
    ],
)
P("Yükleme süresi analizi: Yükleme süresi, program boyutuyla (W word) doğrusal "
  "olarak artar. 115200 baud 8N1'de bir bayt 10 bit = 86.8 µs sürer. W word'ün C=64 "
  "word'lük paketlerle yüklenmesinde host->FPGA bayt sayısı ~ 10*ceil(W/C) + 4*W "
  "olur (paket başı 10 bayt çerçeve + 4W veri). Aşağıdaki tablo ölçülen test "
  "programları ve sentetik büyük programlar için hesaplanan süreleri verir:")
TABLE(
    ["Program boyutu", "Bayt (host->FPGA)", "Hesaplanan süre @115200", "Verim", "Ölçülen süre"],
    [
        ["11 word (test1)", "54", "~4.7 ms", "~9.4 KB/s", "________"],
        ["14 word (test2)", "66", "~5.7 ms", "~9.6 KB/s", "________"],
        ["38 word (test3)", "162", "~14.1 ms", "~10.5 KB/s", "________"],
        ["128 word", "532", "~46.2 ms", "~10.8 KB/s", "________"],
        ["512 word", "2128", "~184.7 ms", "~10.8 KB/s", "________"],
        ["2048 word (8 KB)", "8512", "~738.8 ms", "~10.8 KB/s", "________"],
    ],
)
TODO("'Ölçülen süre' sütununu kart üzerinde doldurun: loader_host.py her yüklemede "
     "gerçek ms değerini ekrana yazdırır. Ayrıca bu tablodan bir GRAFİK (program "
     "boyutu - yükleme süresi) çizip 'Şekil 3.1' olarak ekleyin (PÇ7 'grafik veya tablo').")
P("Eğilim: Süre O(W) — kod büyüdükçe doğrusal artar; çerçeve yükü (paket başı 10 bayt) "
  "büyük programlarda ihmal edilir hâle gelir ve verim ~10.8 KB/s'ye doğru doygunlaşır.")
P("FPGA kaynak tüketimi (Gowin GW1NR-9; toplam 8640 LUT4 / 6480 register / 26 BSRAM). "
  "Aşağıdaki değerler TAHMİNİDİR:")
TABLE(
    ["Modül", "LUT (tahmini)", "Register (tahmini)", "BSRAM", "LUT (gerçek)", "Reg (gerçek)"],
    [
        ["PicoRV32 (MUL/DIV kapalı)", "~1200", "~800", "0", "______", "______"],
        ["uart_rx + uart_tx", "~120", "~90", "0", "______", "______"],
        ["loader FSM (CRC dâhil)", "~150", "~120", "0", "______", "______"],
        ["bram_mem (8 KB)", "~10", "~5", "4", "______", "______"],
        ["TOPLAM (tahmini)", "~1500 (%17)", "~1015 (%16)", "4 / 26", "______", "______"],
    ],
)
TODO("'Gerçek' sütunlarını Gowin EDA sentez raporundan doldurun "
     "(Process -> Place & Route -> Resource Usage Summary). Tahmini değerleri gerçek "
     "değerlerle değiştirip bu uyarıyı silin.")

# ════════════════════════════════════════════════════════════════════════
# BÖLÜM 4
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
H("BÖLÜM 4. PROJENİN KÜRESEL, TOPLUMSAL VE EKONOMİK ETKİLERİ (PÇ8)", size=13)
P("Geliştirilen UART tabanlı loader ve onun üzerine kurulduğu açık kaynaklı RISC-V "
  "araç zincirinin teknik değerinin yanı sıra; Birleşmiş Milletler Sürdürülebilir "
  "Kalkınma Amaçları (SKA) çerçevesinde küresel, toplumsal ve ekonomik etkileri "
  "aşağıda ele alınmıştır.")

H("4.1. Sürdürülebilirlik ve Yeşil Bilişim (SKA 7, SKA 13)", size=11)
P("Loader mekanizmasının en doğrudan çevresel etkisi, GELİŞTİRME DÖNGÜSÜNDEKİ ENERJİ "
  "TASARRUFUDUR. Proje 2'deki sentez-zamanı yükleme yaklaşımında her kod değişikliği, "
  "dakikalar süren ve yüksek işlemci gücü tüketen tam bir FPGA sentez + yer-yerleştirme "
  "(place & route) süreci gerektiriyordu. Loader ile aynı değişiklik, yeniden sentez "
  "olmadan saniyeler içinde (~5-740 ms) UART üzerinden yüklenir. Gün içinde yüzlerce "
  "yineleme yapan bir geliştirici için bu, bilgisayarın tam yük altında çalıştığı "
  "sürenin ve dolayısıyla karbon ayak izinin önemli ölçüde azalması demektir (SKA 13). "
  "Ayrıca RV32I gibi sade bir komut setinin transistör/enerji tasarrufu, yazılım "
  "optimizasyonunun (küçük, verimli makine kodu) enerji tüketimini düşürmedeki rolünü "
  "örnekler (SKA 7).")

H("4.2. Ekonomik Sürdürülebilirlik ve Teknolojik Bağımsızlık (SKA 8, SKA 9)", size=11)
P("Sistemin tamamı açık kaynak bileşenlerden (RISC-V ISA, PicoRV32, Python, Icarus "
  "Verilog) oluşur ve hiçbir ticari lisans gerektirmez. Bu; özellikle gelişmekte olan "
  "ülkeler ve KOBİ'ler için tescilli (kapalı kaynak) işlemci mimarilerine ve pahalı "
  "EDA araçlarına olan bağımlılığı azaltır; yerli çip ve gömülü sistem ekosisteminin "
  "gelişmesini destekler (SKA 9). Düşük maliyetli Tang Nano 9K (~15-20 USD) üzerinde "
  "ücretsiz araç zinciriyle çalışmak, araştırma ve eğitim Ar-Ge maliyetlerini "
  "ciddi ölçüde düşürür ve nitelikli iş gücü yetiştirir (SKA 8).")

H("4.3. Fonksiyonel Güvenlik ve Sağlık (SKA 3)", size=11)
P("Geliştirilen loader'ın CRC-16 tabanlı hata kontrol mekanizması; kritik gömülü "
  "sistemlerde (tıbbi cihazlar, otomotiv, savunma) doğrudan can ve iş güvenliği "
  "boyutuna sahiptir. Bir insülin pompasına veya kalp pili programlayıcısına yüklenen "
  "firmware'deki tek bir bozuk bayt, ölümcül sonuçlara yol açabilir. Bu projede "
  "gösterildiği gibi, bozuk paketin CRC ile yakalanıp NACK ile reddedilmesi ve "
  "yeniden gönderilmesi, 'sessiz veri bozulmasına' (silent data corruption) karşı "
  "temel bir savunma katmanıdır (SKA 3 - Sağlıklı Bireyler).")

H("4.4. E-Atık Yönetimi ve Döngüsel Ekonomi (SKA 12)", size=11)
P("FPGA'in yeniden yapılandırılabilir (reconfigurable) doğası ve bu loader'ın "
  "sağladığı uzaktan/canlı güncelleme (OTA benzeri) yeteneği, tek bir donanımın "
  "yazılım güncellemeleriyle uzun yıllar kullanımda kalmasını sağlar. Aynı kart, "
  "atılmadan farklı programlarla yeniden görevlendirilebilir; bu da elektronik "
  "atığı (e-atık) azaltır ve döngüsel ekonomiyi destekler (SKA 12 - Sorumlu Üretim "
  "ve Tüketim). Sabit-fonksiyonlu (ASIC) bir çözümün aksine, hata düzeltmeleri "
  "donanım değişimi gerektirmez.")

# ════════════════════════════════════════════════════════════════════════
# BÖLÜM 5
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
H("BÖLÜM 5. PROJE YÖNETİMİ VE TAKIM ÇALIŞMASI (PÇ12, PÇ13)", size=13)

H("5.1. Görev Dağılımı ve Sorumluluk Matrisi (RACI)", size=11)
P("Devasa bir sistem entegrasyonu, alt modüllere bölünerek ve her üyeye bağımsız "
  "sorumluluk atanarak yönetilmiştir. Aşağıdaki RACI matrisi (R=Sorumlu, A=Onaylayan, "
  "C=Danışılan, I=Bilgilendirilen) iş bölümünü özetler.")
TABLE(
    ["İş paketi", "A.Hoylu", "F.Özden", "E.Duran", "M.E.Kendir"],
    [
        ["UART RX/TX (Verilog)", "R", "C", "I", "A"],
        ["Loader FSM + CRC (Verilog)", "C", "I", "R", "A"],
        ["top_loader entegrasyon", "A", "C", "C", "R"],
        ["Host gönderici (Python)", "I", "R", "A", "C"],
        ["Test programları (asm)", "R", "A", "C", "I"],
        ["Simülasyon testbench", "C", "I", "A", "R"],
        ["FPGA sentez + kart demo", "A", "R", "C", "C"],
        ["Rapor + video", "R", "R", "R", "R"],
    ],
)
TODO("Yukarıdaki RACI dağılımı ÖRNEKTİR. Gerçek iş bölümünüze göre R/A/C/I "
     "harflerini güncelleyin. İsterseniz ek olarak bir Gantt şeması (zaman çizelgesi) "
     "ile haftalık ilerlemeyi de gösterebilirsiniz (isteğe bağlı, PÇ12'yi güçlendirir).")

H("5.2. Koordinasyon ve Sürüm Kontrol Yönetimi", size=11)
P("Donanım (Verilog) ve yazılım (Python) modüllerinin uyumlu çalışabilmesi için iki "
  "tarafın da paylaştığı sözleşmeler (paket protokolü ve CRC algoritması) önce ortak "
  "kararlaştırılmış, sonra bağımsız gerçeklenmiştir. Entegrasyon riski, donanım "
  "gerektirmeyen Icarus Verilog testbench'i sayesinde erken aşamada giderilmiştir: "
  "host modeli (testbench görevleri) ile gerçek RTL aynı protokolü konuşup "
  "konuşmadığı otomatik olarak doğrulanmış, kart üzerinde bring-up öncesi tüm mantık "
  "hataları ayıklanmıştır. Sürüm kontrolü için Git deposu (PicoRV_Based_Assembler) ve "
  "modül-bazlı dosya organizasyonu (loader/rtl, loader/host, loader/sim) "
  "kullanılmıştır.")

# ════════════════════════════════════════════════════════════════════════
# BÖLÜM 6
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
H("BÖLÜM 6. BİREYSEL KATKI BEYANI (PÇ12)", size=13)
P("Bu bölüm her üye tarafından ayrı ayrı doldurulacak ve ıslak (el) imza ile "
  "imzalanacaktır. Aşağıdaki soru detaylıca cevaplandırılmalıdır:")
P("\"Grup çalışmasından bağımsız olarak, projenin hangi spesifik modüllerini (örn: "
  "UART alıcı, Python gönderici betiği, Loader FSM) tamamen kendi başıma tasarladım? "
  "Bu süreçte karşılaştığım ve tek başıma çözdüğüm en büyük teknik problem neydi?\"", bold=True)
TODO("Bu bölüm AYRI PUANLANIR ve boş bırakılamaz. Her üye kendi paragrafını yazıp "
     "imzalamalı. Aşağıda her üyenin RACI'deki sorumlu (R) olduğu modüllere göre "
     "doldurma rehberi verilmiştir — bu rehberi kendi cümlelerinizle genişletin.")

guide = {
    "ALİHAN HOYLU": "UART RX/TX modülleri ve test assembly programları. Olası teknik "
                    "problem örneği: UART alıcının orta-bit örnekleme zamanlamasını "
                    "(CLKS_PER_BIT) saat kaymasına dayanıklı kuracak şekilde ayarlamak.",
    "FEYZA ÖZDEN":  "Host gönderici (loader_host.py) ve FPGA kart demosu. Olası "
                    "problem örneği: NACK/zaman aşımında paketi güvenilir biçimde "
                    "yeniden gönderen retransmit mantığını kurmak.",
    "ENES DURAN":   "Loader FSM + CRC ve simülasyon testbench. Olası problem örneği: "
                    "CRC-16'nın host (Python) ile FPGA (Verilog) tarafında BİREBİR "
                    "aynı sonucu üretmesini sağlamak.",
    "MUHAMMED EREN KENDİR": "top_loader entegrasyonu ve BRAM/UART port paylaşımı. "
                    "Olası problem örneği: RUN-ACK gönderimi bitmeden CPU reset'ini "
                    "bırakmamak (UART TX hattını loader ile CPU arasında doğru anda "
                    "devretmek).",
}
for m in ["ALİHAN HOYLU", "FEYZA ÖZDEN", "ENES DURAN", "MUHAMMED EREN KENDİR"]:
    P("")
    P(f"Üye: {m}", bold=True)
    gp = doc.add_paragraph()
    gr = gp.add_run(f"(Doldurma rehberi: {guide[m]})")
    _set_font(gr, size=9, color=(0x70, 0x70, 0x70))
    P("Tasarladığım modül(ler): ______________________________________________________")
    P("Karşılaştığım en büyük teknik problem ve çözümü: ______________________________")
    P("____________________________________________________________________________")
    P("İmza: ____________________")

# ════════════════════════════════════════════════════════════════════════
# BÖLÜM 7
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
H("BÖLÜM 7. KAYNAKÇA", size=13)
refs = [
 "[1] Waterman, A., & Asanović, K. (2019). The RISC-V Instruction Set Manual, "
 "Volume I: Unprivileged ISA. RISC-V International. (RV32I komut formatları, "
 "PROGADDR_RESET ve bellek erişimi için referans alınmıştır.)",
 "[2] Beck, L. L. (1997). System Software: An Introduction to Systems Programming "
 "(3rd ed.). Addison-Wesley. (Loader/yükleyici kavramları ve mutlak (absolute) "
 "yükleme algoritmalarının teorik temeli için kullanılmıştır.)",
 "[3] Koopman, P., & Chakravarty, T. (2004). Cyclic Redundancy Code (CRC) Polynomial "
 "Selection for Embedded Networks. IEEE/IFIP DSN 2004. (CRC polinom seçimi ve hata "
 "yakalama gücü analizi için temel alınmıştır.)",
 "[4] Wolf, C. (2019). PicoRV32 - A Size-Optimized RISC-V CPU. YosysHQ GitHub "
 "Deposu. https://github.com/YosysHQ/picorv32 (Hedef yumuşak çekirdek, bellek veri "
 "yolu el sıkışması ve STACKADDR/PROGADDR parametreleri için incelenmiştir.)",
 "[5] STMicroelectronics. (2021). AN3155: USART Protocol Used in the STM32 Bootloader. "
 "Application Note. (UART tabanlı firmware yükleme protokollerinin endüstriyel örneği "
 "olarak incelenmiştir.)",
 "[6] ITU-T Recommendation V.41 (1988). Code-Independent Error-Control System. "
 "International Telecommunication Union. (CRC-16/CCITT (0x1021) standardının tanımı "
 "için kullanılmıştır.)",
 "[7] Python Software Foundation & pySerial. (2023). pySerial Documentation. "
 "https://pyserial.readthedocs.io (Host tarafı seri port iletişimi için referans.)",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(r)
    _set_font(run)

# ════════════════════════════════════════════════════════════════════════
# EK: TESLİM KONTROL LİSTESİ (tüm yapılacaklar tek yerde)
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
H("EK A. TESLİM KONTROL LİSTESİ (YAPILMASI GEREKENLER)", size=13)
P("Aşağıdaki maddeler teslim öncesi tamamlanmalıdır. Tamamlananları [x] ile "
  "işaretleyin. (Kod ve simülasyon tarafı HAZIR ve doğrulanmıştır.)")
TABLE(
    ["#", "Yapılacak", "Sorumlu / Nerede", "Durum"],
    [
        ["1", "Gowin EDA'da top_loader sentezi + bitstream", "FPGA sorumlusu", "[ ]"],
        ["2", ".cst UART pinlerini (17/18) karta göre doğrula", "FPGA sorumlusu", "[ ]"],
        ["3", "3 testi kart üzerinde UART ile çalıştır (LED + sonuç)", "Tüm ekip", "[ ]"],
        ["4", "Bölüm 3.2: gerçek LUT/Register/BSRAM değerleri", "Sentez raporundan", "[ ]"],
        ["5", "Bölüm 3.2: ölçülen yükleme süreleri + grafik", "loader_host.py çıktısı", "[ ]"],
        ["6", "Bölüm 2: blok şema + FSM diyagramı (şekil)", "draw.io/Word", "[ ]"],
        ["7", "Bölüm 5.1: RACI'yi gerçek dağılımla güncelle", "Tüm ekip", "[ ]"],
        ["8", "Bölüm 6: bireysel katkı beyanları + imzalar", "Her üye ayrı", "[ ]"],
        ["9", "Tüm sarı (EKSİK) kutuları kaldır", "Rapor sorumlusu", "[ ]"],
        ["10", "PDF'e çevir + resmi adla kaydet (BIL302_PROJE3_...)", "Rapor sorumlusu", "[ ]"],
        ["11", "Demo videosu (max 5 dk) + resmi adla MP4", "Tüm ekip", "[ ]"],
        ["12", "lms.subu.edu.tr'ye yükle (son tarih 07.06.2026 23:59)", "Bir üye", "[ ]"],
    ],
)

# ── Kaydet ──────────────────────────────────────────────────────────────
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
out = os.path.join(ROOT, "rapor3_taslak.docx")
doc.save(out)
print("Rapor uretildi:", out)
