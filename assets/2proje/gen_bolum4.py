from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FONT = "Courier New"
SZ = 10

def p(doc, text="", bold=False, italic=False, sz=SZ):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(sz)
    run.bold = bold
    run.italic = italic
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(1 if sz < SZ else 2)
    return para

def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(SZ)
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = str(val)
            for run in c.paragraphs[0].runs:
                run.font.name = FONT
                run.font.size = Pt(SZ)
    return t

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2)
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(SZ)

# ── BAŞLIK ──────────────────────────────────────────────────────────────
p(doc, "BÖLÜM 4. OBJECT DOSYA FORMATI (PÇ 6)", bold=True)
p(doc, "─" * 72)
p(doc)
p(doc,
  "Assembler'ın ürettiği makine kodu ve sembol bilgileri, linker tarafından "
  "tüketilebilecek yapılandırılmış bir ara formata dönüştürülmektedir. Bu "
  "formata ObjectFile adı verilmekte; hem bellekte Python nesnesi olarak "
  "tutulmakta hem de .o uzantılı JSON dosyasına serileştirilebilmektedir.")
p(doc)

# 4.1
p(doc, "4.1. ObjectFile Veri Yapısı", bold=True)
p(doc)
p(doc,
  "Her .asm kaynak dosyası bağımsız olarak derlenerek bir ObjectFile nesnesi "
  "üretir. Nesnenin alanları aşağıdaki tabloda tanımlanmıştır:")
p(doc)

table(doc,
  ["Alan", "Tip", "Açıklama"],
  [
    ["name",        "str",  "Modül adı (linker haritasında kullanılır)"],
    ["text",        "list", "[(adres, değer, boyut)]  — .text section kelimeleri"],
    ["data",        "list", "[(adres, değer, boyut)]  — .data section kelimeleri"],
    ["globals",     "dict", "{etiket: adres}  — dışarıya açık semboller"],
    ["externs",     "list", "[etiket]  — dışarıdan beklenen semboller"],
    ["relocations", "list", "[(adres, etiket, tip)]  — çözümsüz referanslar"],
    ["symtab",      "dict", "{etiket: adres}  — tüm lokal semboller"],
  ]
)
p(doc)
p(doc,
  "text ve data ayrımı, kaynak dosyadaki .data direktifinin bulunduğu adres "
  "esas alınarak from_assembler() fabrika metodu içinde yapılmaktadır. "
  "Bu sayede linker, .text ve .data section'larını bağımsız adres alanlarına "
  "yerleştirebilmektedir.")
p(doc)

# 4.2
p(doc, "4.2. H/D/R/T/M/E Kayıt Formatı", bold=True)
p(doc)
p(doc,
  "GUI'nin Object Dosyaları sekmesinde ve get_object_record() metoduyla "
  "üretilen metin çıktısı, SIC/XE standardından esinlenerek tasarlanmış "
  "altı kayıt tipinden oluşmaktadır:")
p(doc)

table(doc,
  ["Kayıt", "Sembol", "İçerik"],
  [
    ["Header",        "H", "H<modül_adı:6><başlangıç:6X><uzunluk:6X>"],
    ["Define",        "D", "D<etiket:6><adres:6X> [...]  — global semboller"],
    ["Refer",         "R", "R<etiket:6> [...]  — extern semboller"],
    ["Text",          "T", "T<başlangıç:6X><byte_sayısı:2X><hex_kod>"],
    ["Modification",  "M", "M<reloc_adresi:6X>05+<etiket:6>"],
    ["End",           "E", "E<giriş_noktası:6X>"],
  ]
)
p(doc)
p(doc,
  "T kaydında ardışık kelimeler tek satırda birleştirilir; adres sürekliliği "
  "kırıldığında yeni bir T kaydı başlatılır. M kaydındaki '05' alanı, "
  "değiştirilecek alanın yarı-byte (nibble) genişliğini belirtir ve "
  "SIC/XE convention'ını izler.")
p(doc)

# 4.2 — Gerçek örnekler
p(doc, "4.2.1. Örnek: math.asm Object Kaydı", bold=True)
p(doc)
p(doc,
  "Fibonacci yardımcı modülü (FIB_STEP fonksiyonu, 2 komut = 8 byte):")
p(doc)
for line in [
    "HMATH  000200000008",
    "DFIB_STEP000200",
    "T0002000800B5053300008067",
    "E000200",
]:
    p(doc, f"  {line}", sz=8)
p(doc)

table(doc,
  ["Satır", "Açıklama"],
  [
    ["HMATH  000200000008",       "Modül: MATH, başlangıç: 0x000200, uzunluk: 8 byte"],
    ["DFIB_STEP000200",           "Global sembol FIB_STEP, adres: 0x000200"],
    ["T0002000800B5053300008067", "0x200 adresinden 8 byte: add+jalr makine kodları"],
    ["E000200",                   "Giriş noktası: 0x000200 (FIB_STEP)"],
  ]
)
p(doc)

p(doc, "4.2.2. Örnek: main.asm Object Kaydı (Fibonacci)", bold=True)
p(doc)
p(doc,
  "Ana modül (MAIN, .text + .data bölümlü, FIB_STEP extern referansı):")
p(doc)
for line in [
    "HMAIN  000000010028",
    "DMAIN  000000",
    "RFIB_STEP",
    "T0000001C000005130010059300A00293...",
    "T01000028...",
    "M00004005+FIB_STEP",
    "E000000",
]:
    p(doc, f"  {line}", sz=8)
p(doc)

table(doc,
  ["Satır", "Açıklama"],
  [
    ["HMAIN  000000010028", "Modül: MAIN, başlangıç: 0x000000, toplam alan: 0x10028 byte"],
    ["DMAIN  000000",       "Global sembol MAIN, adres: 0x000000"],
    ["RFIB_STEP",           "Extern referans: FIB_STEP (henüz çözümlenmemiş)"],
    ["T0000001C ...",       ".text section — 0x000000 adresinden itibaren makine kodu"],
    ["T010000 ...",         ".data section — 0x010000 adresinden itibaren veri kelimeleri"],
    ["M00004005+FIB_STEP",  "0x000040 adresindeki JAL komutu FIB_STEP adresiyle yamanacak"],
    ["E000000",             "Giriş noktası: 0x000000 (MAIN)"],
  ]
)
p(doc)

# 4.3
p(doc, "4.3. JSON Serileştirme (.o Formatı)", bold=True)
p(doc)
p(doc,
  "ObjectFile nesneleri save() metodu ile .o uzantılı JSON dosyasına "
  "yazılabilir, load() metodu ile geri yüklenebilir. Bu yaklaşım, önceden "
  "derlenip saklanan modül kütüphanelerinin tekrar derleme yapılmaksızın "
  "linklenmesine olanak tanımaktadır.")
p(doc)
p(doc, "math.o içeriği (özet):")
p(doc)
for line in [
    '{',
    '  "name": "math",',
    '  "text": [',
    '    [512, 11863347, 4],   // 0x200: add x10,x10,x11  => 0x00B50533',
    '    [516, 32871,    4]    // 0x204: jalr x0,x1,0      => 0x00008067',
    '  ],',
    '  "data": [],',
    '  "globals":     { "FIB_STEP": 512 },',
    '  "externs":     [],',
    '  "relocations": [],',
    '  "symtab":      { "FIB_STEP": 512 }',
    '}',
]:
    p(doc, f"  {line}", sz=8)
p(doc)
p(doc,
  "Adres değerleri ondalık (Python int) olarak saklanır; H/D/T/M/E formatına "
  "dönüşüm get_object_record() çağrısı sırasında gerçekleştirilir.")
p(doc)

# kaydet
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bolum4.docx")
doc.save(out)
print(f"Kaydedildi: {out}")
