from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FONT = "Courier New"
SZ = 10

def add_para(doc, text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SZ)
    run.bold = bold
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(SZ)
            run.bold = True
    # data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = FONT
                run.font.size = Pt(SZ)
    return table

doc = Document()

# margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# default font
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(SZ)

# ── BAŞLIK ──────────────────────────────────────────────────────────────
add_para(doc, "BÖLÜM 12. ALGORİTMA KARMAŞIKLIĞI ANALİZİ (PC 6)", bold=True)
add_para(doc, "─" * 72)
add_para(doc)

# 12.1
add_para(doc, "12.1. Ölçüm Yöntemi ve Metodoloji", bold=True)
add_para(doc)
add_para(doc,
    "Linker'in performans analizi için Python'un time.perf_counter() modülü "
    "kullanılmıştır. Dört bağımsız deney tasarlanmış; her deney 5 kez "
    "tekrarlanarak ortalama (ort), minimum (min), maksimum (max) ve standart "
    "sapma (std) değerleri hesaplanmıştır. Testler "
    "tests/test_linker_complexity.py dosyasında yer almaktadır.")
add_para(doc)
add_para(doc, "  Deney 1 : Tek dosya, artan komut sayısı      (N = 10 … 2000)")
add_para(doc, "  Deney 2 : Artan dosya sayısı, sabit komut    (2 … 20 dosya x 10 komut)")
add_para(doc, "  Deney 3 : Artan global sembol sayısı         (S = 5 … 200)")
add_para(doc, "  Deney 4 : Artan relocation kaydı sayısı      (R = 1 … 20)")
add_para(doc)

# 12.2
add_para(doc, "12.2. Ampirik Ölçüm Sonuçları", bold=True)
add_para(doc)

# Deney 1
add_para(doc, "Deney 1 — Tek Dosya, Artan Komut Sayısı", bold=True)
add_table(doc,
    ["N (komut)", "ort (ms)", "min (ms)", "max (ms)", "std (ms)"],
    [
        [10,    "0.091", "0.057",  "0.201",  "0.062"],
        [50,    "0.231", "0.221",  "0.252",  "0.012"],
        [100,   "0.448", "0.436",  "0.462",  "0.010"],
        [200,   "0.913", "0.874",  "1.009",  "0.057"],
        [500,   "2.185", "2.114",  "2.272",  "0.063"],
        [1000,  "4.663", "4.551",  "4.782",  "0.098"],
        [2000,  "9.407", "8.941", "10.106",  "0.533"],
    ]
)
add_para(doc, "  N oranı: 200x  |  Süre oranı: 103.0x  =>  O(N) ONAYLANDI")
add_para(doc)

# Deney 2
add_para(doc, "Deney 2 — Artan Dosya Sayısı (Sabit Komut/Dosya)", bold=True)
add_table(doc,
    ["Dosya", "Toplam N", "ort (ms)", "min (ms)", "max (ms)", "std (ms)"],
    [
        [2,  20,  "0.111", "0.101", "0.140", "0.016"],
        [3,  30,  "0.151", "0.146", "0.161", "0.006"],
        [5,  50,  "0.249", "0.241", "0.259", "0.008"],
        [8,  80,  "0.398", "0.388", "0.403", "0.007"],
        [10, 100, "0.526", "0.494", "0.630", "0.059"],
        [15, 150, "0.764", "0.729", "0.816", "0.033"],
        [20, 200, "0.990", "0.974", "1.017", "0.019"],
    ]
)
add_para(doc, "  Dosya oranı: 10x  |  Süre oranı: 8.9x  =>  O(N) ONAYLANDI")
add_para(doc)

# Deney 3
add_para(doc, "Deney 3 — Artan Global Sembol Sayısı", bold=True)
add_table(doc,
    ["S (sembol)", "ort (ms)", "min (ms)", "max (ms)", "std (ms)"],
    [
        [5,   "0.042", "0.038", "0.051", "0.005"],
        [10,  "0.082", "0.069", "0.120", "0.022"],
        [20,  "0.159", "0.132", "0.264", "0.058"],
        [50,  "0.332", "0.325", "0.349", "0.010"],
        [100, "0.719", "0.674", "0.814", "0.056"],
        [200, "1.562", "1.481", "1.624", "0.066"],
    ]
)
add_para(doc, "  Sembol oranı: 40x  |  Süre oranı: 37.1x  =>  O(S) ONAYLANDI")
add_para(doc)

# Deney 4
add_para(doc, "Deney 4 — Artan Relocation Kaydı Sayısı", bold=True)
add_table(doc,
    ["R (relocation)", "ort (ms)", "min (ms)", "max (ms)", "std (ms)"],
    [
        [1,  "0.044", "0.035", "0.057", "0.009"],
        [2,  "0.064", "0.058", "0.072", "0.005"],
        [5,  "0.121", "0.117", "0.128", "0.004"],
        [10, "0.228", "0.221", "0.235", "0.006"],
        [20, "0.486", "0.429", "0.660", "0.098"],
    ]
)
add_para(doc, "  Relocation oranı: 20x  |  Süre oranı: 11.0x  =>  O(R) ONAYLANDI")
add_para(doc)

# 12.3
add_para(doc, "12.3. Teorik Zaman Karmaşıklığı Analizi", bold=True)
add_para(doc)
add_table(doc,
    ["Aşama", "Açıklama", "Karmaşıklık"],
    [
        ["Section Layout (_layout_section)", "Her kelime tek geçişte işlenir",          "O(N)"],
        ["Sembol birleştirme",               "Her global sembol bir kez eklenir",        "O(S)"],
        ["word_index dict oluşturma",        "linked_code üzerinde tek geçiş",           "O(N)"],
        ["Relocation yamalama",              "Her kayıt için O(1) dict erişimi",         "O(R)"],
        ["Genel linker",                     "",                                          "O(N+S+R)"],
    ]
)
add_para(doc)
add_para(doc,
    "N, S ve R programın büyüklüğüyle doğrusal arttığı için linker toplamda "
    "O(N) zaman karmaşıklığına sahiptir. Ampirik ölçümler bu teorik analizi "
    "doğrulamaktadır: girdi boyutu 200x artırıldığında süre yalnızca 103x "
    "artmış, beklenen O(N) sınırının altında kalmıştır.")
add_para(doc)

# 12.4
add_para(doc, "12.4. Alan Karmaşıklığı Analizi", bold=True)
add_para(doc)
add_table(doc,
    ["Yapı", "İçerik", "Karmaşıklık"],
    [
        ["linked_code",   "Tüm kelimeler bellekte tutulur", "O(N)"],
        ["global_symtab", "Tüm global semboller",           "O(S)"],
        ["word_index",    "Adres→indeks eşlemesi",          "O(N)"],
        ["link_map",      "Her kelime için bir kayıt",      "O(N)"],
        ["Genel",         "",                               "O(N+S)"],
    ]
)
add_para(doc)

# ── KAYDET ──────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bolum12.docx")
doc.save(out)
print(f"Kaydedildi: {out}")
