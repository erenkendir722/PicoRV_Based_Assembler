from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FONT = "Courier New"
SZ = 10

def add_para(doc, text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SZ)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(doc, headers, rows, pass_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = FONT
                run.font.size = Pt(SZ)
                run.bold = True
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT
                    run.font.size = Pt(SZ)
                    # pass_col: o sütundaki GEÇTİ/KALDI renklendirmesi
                    if pass_col is not None and ci == pass_col:
                        if "GEÇTİ" in str(val):
                            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                        elif "KALDI" in str(val):
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    return table

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(SZ)

# ── BAŞLIK ──────────────────────────────────────────────────────────────
add_para(doc, "BÖLÜM 11. TEST SENARYOLARI VE SONUÇLAR (PÇ 7)", bold=True)
add_para(doc, "─" * 72)
add_para(doc)
add_para(doc,
    "Geliştirilen linker yazılımının doğruluğunu ve güvenilirliğini kanıtlamak "
    "amacıyla üç katmanlı bir test stratejisi uygulanmıştır. Testler "
    "tests/test_linker.py dosyasında yer almakta olup 2026-05-08 tarihinde "
    "çalıştırılmıştır.")
add_para(doc)

# 11.1
add_para(doc, "11.1. Birim Testleri", bold=True)
add_para(doc)
add_para(doc,
    "Linker'in temel bileşenleri izole edilerek aşağıdaki birim testleri "
    "uygulanmıştır:")
add_para(doc)

add_table(doc,
    ["#", "Test Adı", "Açıklama", "Sonuç"],
    [
        ["1",  "Tek dosya assemble",         "Tek ObjectFile derleme başarısı",                "GEÇTİ ✓"],
        ["2",  "Tek dosya link",             "Tek modülün linker'dan geçmesi",                 "GEÇTİ ✓"],
        ["3",  "linked_code boş değil",      "Çıktı listesinin dolu olması",                   "GEÇTİ ✓"],
        ["4",  "text_base 0x0 doğru",        "İlk adresin 0x00000000 olması",                  "GEÇTİ ✓"],
        ["5",  "MAIN global_symtab'ta",      "Global sembol tablosuna kayıt",                  "GEÇTİ ✓"],
        ["6",  "MAIN adresi 0x0",            "MAIN'in doğru adrese atanması",                  "GEÇTİ ✓"],
        ["7",  "Çakışan global sembol → hata","Aynı etiket iki modülde → LinkerError",         "GEÇTİ ✓"],
        ["8",  "Hata mesajı 'DUP' içeriyor", "Hata metninin sembol adını içermesi",            "GEÇTİ ✓"],
        ["9",  "Çözümsüz extern → link hatası","Tanımsız extern sembol → hata",               "GEÇTİ ✓"],
        ["10", "Hata mesajı 'HAYALET' içeriyor","Hata metninin sembol adını içermesi",         "GEÇTİ ✓"],
        ["11", "Özel base adresiyle link",   "text_base=0x1000, data_base=0x5000",             "GEÇTİ ✓"],
        ["12", "text 0x1000'den başlıyor",   ".text section adres doğrulaması",                "GEÇTİ ✓"],
        ["13", "data 0x5000'den başlıyor",   ".data section adres doğrulaması",                "GEÇTİ ✓"],
        ["14", "DAT[0] = 0xDEAD",            "Veri kelimesinin doğru yüklenmesi",              "GEÇTİ ✓"],
        ["15", "DAT[1] = 0xBEEF",            "Ardışık veri kelimesinin doğruluğu",             "GEÇTİ ✓"],
    ],
    pass_col=3
)
add_para(doc)

# 11.2
add_para(doc, "11.2. Entegrasyon Testleri", bold=True)
add_para(doc)
add_para(doc,
    "Birden fazla modülün birlikte çalıştığı gerçek senaryolar test edilmiştir:")
add_para(doc)

add_table(doc,
    ["#", "Senaryo", "Modüller", "Doğrulanan Özellik", "Sonuç"],
    [
        ["1",  "Çoklu dosya extern çözümleme",
               "main.asm + foo.asm",
               "JAL opcode=0x6F, FOO adresi>0, extern listesi dolu",
               "GEÇTİ ✓"],
        ["2",  "Çözümsüz extern → assemble tamam",
               "c.asm (HAYALET extern)",
               "Assemble başarılı, link aşamasında hata üretilmesi",
               "GEÇTİ ✓"],
        ["3",  "Fibonacci entegrasyon",
               "fibonacci/main.asm + math.asm",
               "FIB_STEP>MAIN adresi, .mem/.hex çıktısı dolu, link_map dolu",
               "GEÇTİ ✓"],
        ["4",  "main extern FOO içeriyor",
               "main.asm",
               "ObjectFile.externs listesinin doğru doldurulması",
               "GEÇTİ ✓"],
        ["5",  "foo global FOO içeriyor",
               "foo.asm",
               "ObjectFile.globals listesinin doğru doldurulması",
               "GEÇTİ ✓"],
        ["6",  "fibonacci/main.asm assemble",
               "fibonacci/main.asm",
               "Gerçek proje dosyasının hatasız derlenmesi",
               "GEÇTİ ✓"],
        ["7",  "fibonacci/math.asm assemble",
               "fibonacci/math.asm",
               "FIB_STEP fonksiyonunun hatasız derlenmesi",
               "GEÇTİ ✓"],
        ["8",  "Fibonacci link başarılı",
               "main + math",
               "İki modülün birleştirilmesi ve relocation tamamlanması",
               "GEÇTİ ✓"],
        ["9",  ".mem çıktısı boş değil",
               "main + math",
               "FPGA için $readmemh formatı üretilmesi",
               "GEÇTİ ✓"],
        ["10", ".hex çıktısı boş değil",
               "main + math",
               "Intel HEX formatı üretilmesi",
               "GEÇTİ ✓"],
        ["11", "link_map dolu",
               "main + math",
               "Linkleme haritasının üretilmesi",
               "GEÇTİ ✓"],
    ],
    pass_col=4
)
add_para(doc)

# Özet satırı
add_para(doc, "  TOPLAM: 34  |  GEÇTİ: 34  |  KALDI: 0  (2026-05-08)", bold=True)
add_para(doc)

# 11.3
add_para(doc, "11.3. FPGA Doğrulama", bold=True)
add_para(doc)
add_para(doc,
    "Fibonacci ve Mors kodu programlarının ürettiği .mem çıktıları Tang Nano 9K "
    "kartı üzerindeki PicoRV32 çekirdeğine yüklenmiş; LED çıkışları gözlemlenerek "
    "programların doğru çalıştığı doğrulanmıştır.")
add_para(doc)
add_para(doc, "  [BURAYA: FPGA gösterim videosu / ekran görüntüsü eklenecek]", italic=True)
add_para(doc)

# kaydet
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bolum11.docx")
doc.save(out)
print(f"Kaydedildi: {out}")
